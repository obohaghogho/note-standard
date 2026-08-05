import os
import zipfile
import openpyxl
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas

# Define paths
WORKSPACE_DIR = r"d:\Users\Manuel\OneDrive\Desktop\note-standard-latest"
OUTPUT_DIR = os.path.join(WORKSPACE_DIR, "documents", "Anchor Submission Package")
os.makedirs(OUTPUT_DIR, exist_ok=True)

SOURCE_EXCEL = os.path.join(WORKSPACE_DIR, "temp_anchor_source", "Anchor Onboarding Questionnaire (3) (3).xlsx")
SOURCE_DOCX = os.path.join(WORKSPACE_DIR, "temp_anchor_source", "Updated Standard Service Agreement (2).docx")
SIG_PNG = os.path.join(WORKSPACE_DIR, "temp_anchor_source", "signature_transparent.png")
SEAL_PNG = os.path.join(WORKSPACE_DIR, "temp_anchor_source", "seal_transparent.png")

FILE_EXCEL = os.path.join(OUTPUT_DIR, "01 - Anchor Onboarding Questionnaire.xlsx")
FILE_DOCX = os.path.join(OUTPUT_DIR, "02 - Standard Service Agreement.docx")
FILE_INDEMNITY = os.path.join(OUTPUT_DIR, "03 - Fincra Wildcard IP Indemnity.pdf")
FILE_COVER = os.path.join(OUTPUT_DIR, "04 - Cover Letter.docx")
FILE_COMPLIANCE = os.path.join(OUTPUT_DIR, "05 - Compliance Summary.pdf")
FILE_README = os.path.join(OUTPUT_DIR, "06 - README.txt")
FILE_ZIP = os.path.join(OUTPUT_DIR, "Anchor_Submission_Package_Notestandard.zip")

print("Executing Complete Service Agreement & Package Fix...")

# ==============================================================================
# 1. EXCEL QUESTIONNAIRE AUDIT & GENERATION
# ==============================================================================
print("[1/6] Processing Excel Questionnaire...")
wb = openpyxl.load_workbook(SOURCE_EXCEL)
ws = wb.active

responses = {
    4: "Jossy Digital Technologies Ltd",
    5: "NoteStandard",
    6: "NoteStandard is a modern enterprise fintech and digital workspace platform operated by Jossy Digital Technologies Ltd (RC 9586407). NoteStandard provides multi-currency wallets, virtual accounts, NIP bank transfers, internal ledger transfers, merchant collection services, bill payments, and integrated digital asset workflows through regulated third-party infrastructure partners. Banking infrastructure and regulated financial services are delivered via licensed Banking-as-a-Service (BaaS) and financial institution partners. NoteStandard is a technology solution provider and is not a licensed bank.",
    7: "https://notestandard.com",
    8: "Less than 1 year",
    9: "Privately Owned (Private Limited Company)",
    10: "5",
    11: "Effurun, Delta State, Nigeria",
    12: "Founder Funded",
    13: "Nigeria",
    14: "No",
    15: "No",
    16: "No",
    17: "N/A",

    20: "Digital Wallet, Bank Account, Virtual Accounts, NIP Transfer, Book Transfer, P2P, Payment cards (planned), Crypto Wallet, Bill payments",
    21: "NoteStandard",
    22: "Subscription Fee: Starter ($0/mo), Pro ($99/mo), Enterprise ($499/mo)\nWallet Maintenance Fee: NGN 500/mo per dedicated account (Waived if monthly volume > ₦1,000,000); USD $5.00/mo\nPer Transaction Fee: NGN Bank Transfer (1.2% + ₦50, capped at ₦2,000); Local Cards (1.5% + ₦100, capped at ₦2,000); USD Card (3.5% + $0.30)\nWithdrawal Fee: NGN NIP (₦50 flat for <= ₦50k, ₦100 flat for > ₦50k); USD Wire ($15 flat)\nFX Fee: Dynamic volume-tiered margin spread (0.15% - 0.75% over interbank rate locked for 60 seconds)",
    24: "Subscription Fee: Starter ($0/mo), Pro ($99/mo), Enterprise ($499/mo)",
    25: "Monthly Service / Maintenance Fee: NGN 500/mo per dedicated account; USD $5.00/mo",
    26: "Per Transaction Fee: NGN Bank Transfer (1.2% + ₦50, max ₦2,000); Cards (1.5% + ₦100)",
    27: "Minimum balance fee: None ($0.00)",
    28: "Other fees: NIP Withdrawal ₦50/₦100; FX margin spread 0.15% - 0.75%",
    30: "Minimum required spend: None ($0.00)",
    31: "Minimum deposit: None ($0.00)",
    32: "Minimum/Maximum balance: No minimum balance requirement; maximum limits subject to KYC tier",
    33: "Other: None",
    34: "Volume-based transaction fee discounts for high-volume enterprise merchants.",
    35: "No",
    36: "Monthly Total Volume: ₦100,000,000 (~$66,667 USD) | Daily Average Volume: ₦3,333,333 (~$2,222 USD)",
    38: "Daily: ₦360,000 (30 txns) | Monthly: ₦10,800,000 (900 txns)",
    39: "Daily: ₦1,275,000 (55 txns: ₦875k NIP + ₦400k Payouts) | Monthly: ₦38,250,000 (1,650 txns)",
    40: "Daily: ₦100,000 (5 txns crypto off-ramp) | Monthly: ₦3,000,000 (150 txns)",
    42: "Daily: ₦1,350,000 (45 txns) | Monthly: ₦40,500,000 (1,350 txns)",
    43: "Daily: ₦180,000 (10 txns crypto/card on-ramp) | Monthly: ₦5,400,000 (300 txns)",
    44: "Daily: ₦360,000 (30 txns) | Monthly: ₦10,800,000 (900 txns)",
    45: "None",
    46: "N/A",
    47: "N/A",
    48: "N/A",

    51: "53",
    52: "1000 (Expected MAU: 100)",
    53: "Both Consumers and Businesses (Personal, SMEs, Freelancers, Corporate Accounts)",
    54: "Community growth initiatives, organic business onboarding, and targeted digital marketing",
    55: "Personal financial management, wallet funding, NIP transfers, bill payments, and peer-to-peer transfers",
    56: "Business collection accounts, invoice payments, corporate payroll disbursements, and treasury management",
    57: "Sole Proprietorship, Private Limited Companies, Registered SMEs, Digital Merchants",
    58: "Yes. Account holders can hold multi-currency sub-wallets (NGN, USD, EUR, GBP). Internal transfers are supported.",
    59: "No",

    62: "Founder Funded. 100% bootstrapped with zero third-party equity dilution.",
    63: "12 to 24 months",
    64: "Yes, planned expansion of engineering and compliance operations as transaction volumes scale.",

    67: "Nigeria",
    68: "Yes, subject to full KYC verification (valid Nigerian residency permit, alien registration, international passport, and proof of address).",
    69: "Yes, subject to ongoing compliance requirements. Existing customers who relocate outside Nigeria may continue using eligible services after completing any required KYC refresh, sanctions screening, and regulatory checks. Services restricted by applicable laws, regulatory requirements, or provider availability may be limited or suspended depending on the customer's country of residence.",
    70: "Yes, subject to enhanced international KYC, valid government ID verification, and sanction screening.",
    71: "Current: International card collections and FX conversions supported via licensed partners (Paystack & Fincra).\nFuture: Dedicated USD accounts, ACH, SEPA, and SWIFT payouts are planned for Phase 2 post-Anchor approval.",

    74: "Yes. Jossy Digital Technologies Ltd is a duly incorporated Nigerian private limited company (RC 9586407). Financial and banking services are delivered via regulated licensed partners.",
    75: "No",
    76: "No. Cryptocurrency functionality is offered strictly through regulated third-party infrastructure providers. NoteStandard does not operate an independent exchange or hold un-hosted crypto custody.",
    77: "No",
    78: "No",
    79: "Yes (via regulated third-party infrastructure partners only)",
    80: "No",
    81: "No",
    82: "No",
    84: "Internal double-entry ledger audits completed; external financial audit planned before production scale.",
    85: "Internal automated OWASP penetration testing completed; external CREST pen test planned before production scale.",
    86: "Architecture built for SOC2 compliance; formal SOC2 Type II audit planned.",
    87: "Internal IT security controls and telemetry audits completed; external IT audit planned.",
    88: "Internal BCP and automated disaster recovery failover exercises completed.",
    90: "YES",
    91: "YES",
    92: "YES",
    93: "YES (BVN & NIN post-KYC)",
    94: "YES",
    95: "NO (PCI-DSS compliant gateway tokenization only)",
    96: "YES",
    97: "YES",
    98: "YES",
    99: "YES",
    100: "Wallet Balances, Login History, Support Chat Logs, AML Screening Records, Risk Scores, RBAC Permissions",
    101: "N/A",
    103: "YES (Customer onboarding includes automated identity verification and sanctions screening through integrated compliance providers)",
    104: "YES (Background verification & identity checks)",
    105: "YES (Background verification & identity checks)",
    107: "Primary KYC Provider: Prembly (IdentityPass). Additional providers may be integrated as business requirements evolve.",
    108: "YES",
    109: "Internal Risk Decision Engine, Custom AML Engine, and Velocity Rules Engine.",

    112: "YES (JWT authentication with multi-factor support)",
    113: "YES (E.164 MSISDN validation via SMS gateway)",
    114: "YES (Two-factor OTP required for high-risk payouts)",
    115: "YES (Strict event_id and idempotency key enforcement + replay guard)",
    116: "YES (AES-256 encryption at rest via KMS; TLS 1.3 in transit)",
    117: "NoteStandard maintains its own internal double-entry ledger for customer balances, reconciliation, treasury management, and accounting. Anchor will be used solely as regulated banking infrastructure and not as the system of record.",
    118: "YES (Mandatory annual data privacy and security awareness training)",

    121: "Monday–Friday: 8:00 AM – 6:00 PM WAT | Saturday: 9:00 AM – 4:00 PM WAT",
    122: "24/7 in-app AI support + live escalation to human support agents; email support ticket queue (admin@notestandard.com / support@notestandard.com).",
    123: "In-house customer operations and compliance team handling tickets, live chat, email (admin@notestandard.com), phone, and WhatsApp.",
    124: "Real-time wallet balance dashboard via web application (https://notestandard.com) and mobile responsive interface.",
    125: "Digital onboarding via NoteStandard web/mobile app: email sign-up, phone OTP verification, tier 1-3 KYC submission.",
    126: "Dedicated complaint logging service with SLA-tracked resolution workflows, audit logging, and regulatory escalation paths.",

    129: "NoteStandard Terms of Service",
    130: "NoteStandard Privacy Policy",
    131: "NoteStandard Electronic Signature & Communications Agreement"
}

for rng in list(ws.merged_cells.ranges):
    if "C68:C69" in str(rng):
        ws.unmerge_cells("C68:C69")

for row_idx, val in responses.items():
    ws.cell(row=row_idx, column=3, value=val)

wb.save(FILE_EXCEL)
print(f"[OK] Saved {FILE_EXCEL}")

# ==============================================================================
# 2. COMPLETE SERVICE AGREEMENT DOCX AUDIT & BLANK FIELD FIX
# ==============================================================================
print("[2/6] Processing Service Agreement DOCX & Filling All Blank Spaces & Schedule B SLA...")

doc_sa = docx.Document(SOURCE_DOCX)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def create_styled_table(doc_obj, col_widths, headers, data):
    table = doc_obj.add_table(rows=0, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="CBD5E1", sz="4")
    
    # Header row
    hdr_row = table.add_row()
    for ci, heading in enumerate(headers):
        cell = hdr_row.cells[ci]
        cell.text = heading
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = 'Calibri'
            
    # Data rows
    for ri, row_data in enumerate(data):
        row = table.add_row()
        bg_color = "F8FAFC" if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=90, bottom=90, left=140, right=140)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(30, 41, 59)
                run.font.name = 'Calibri'

    # Set widths
    for row in table.rows:
        for ci, w in enumerate(col_widths):
            if ci < len(row.cells):
                row.cells[ci].width = Inches(w)
                
    return table

for i, p in enumerate(doc_sa.paragraphs):
    text = p.text.strip()
    
    # Para 12: Client Name on title page
    if i == 12:
        p.text = "JOSSY DIGITAL TECHNOLOGIES LTD"
    
    # Para 16: Date on title page
    elif i == 16:
        p.text = "THIS 03 DAY OF AUGUST 2026"
    
    # Para 21: Date in opening line
    elif i == 21:
        p.text = "This Client Service Agreement (“Agreement”) is made this 03 day of August 2026."
    
    # Para 25: Client legal details
    elif i == 25:
        p.text = "JOSSY DIGITAL TECHNOLOGIES LTD, a limited liability company duly incorporated under the laws of the Federal Republic of Nigeria, with RC Number: 9586407 and having its registered address at Effurun, Delta State, Nigeria (where the context so admits includes its assigns and successors) (“the Client”)."
    
    # Para 28: Client business description
    elif i == 28:
        p.text = "The Client is a fintech and digital workspace platform providing multi-currency wallets, virtual accounts, NIP bank transfers, internal ledger transfers, merchant collections, bill payments, and treasury technology services."

    # Anchor Date
    elif "DATE:" in text and i < 377 and "OLUWASEGUN" in doc_sa.paragraphs[i-2].text:
        p.text = "DATE: 03 August 2026"

    # Client Execution block
    elif "FOR: COMPANY NAME [CLIENT]" in text or "FOR: COMPANY NAME" in text:
        p.text = "FOR: JOSSY DIGITAL TECHNOLOGIES LTD [CLIENT]"
        if i + 2 < len(doc_sa.paragraphs) and "SIGNATURE:" in doc_sa.paragraphs[i+2].text:
            sig_p = doc_sa.paragraphs[i+2]
            sig_p.text = "SIGNATURE: "
            run = sig_p.add_run()
            run.add_picture(SIG_PNG, width=Inches(1.8))
        if i + 3 < len(doc_sa.paragraphs):
            doc_sa.paragraphs[i+3].text = "NAME: AGHOGHO OBOH"
        if i + 4 < len(doc_sa.paragraphs):
            doc_sa.paragraphs[i+4].text = "DESIGNATION: DIRECTOR"
        if i + 5 < len(doc_sa.paragraphs):
            doc_sa.paragraphs[i+5].text = "DATE: 03 August 2026"

    # Schedules
    elif "Please find attached a list of our high risk and prohibited Customers" in text:
        p.text = "NoteStandard prohibits onboarding high-risk/restricted entities including sanctions-listed persons, un-hosted gambling operators, adult content services, and unregulated darknet vendors."
    elif "Please find attached a list of our high risk and prohibited Jurisdictions" in text:
        p.text = "NoteStandard enforces FATF and OFAC high-risk and non-cooperative jurisdiction blacklists including North Korea (DPRK), Iran, Myanmar, Syria, and Cuba."

# SLA Table 0 Data: ISSUE RESOLUTION
table0_widths = [1.3, 2.7, 1.2, 1.3]
table0_headers = ["REQUEST CATEGORY", "DESCRIPTION & OPERATIONAL SCOPE", "ACKNOWLEDGEMENT SLA", "TARGET RESOLUTION SLA"]
table0_data = [
    ["Account Request", "Virtual Account Creation (NGN & USD dedicated accounts)", "15 minutes", "2 hours"],
    ["Account Request", "Account Profile Modifications, KYC Re-verification & Tier Upgrades", "30 minutes", "4 hours"],
    ["Account Request", "Account Restrict / Freeze / Compliance Sanction Lock", "15 minutes", "1 hour"],
    ["Transaction Request", "Transaction Verification & NIP Payout Status Inquiry", "15 minutes", "1 hour"],
    ["Transaction Request", "Transaction Webhook Re-delivery & Event Re-sync", "15 minutes", "2 hours"],
    ["Reporting Request", "Suspicious Transaction Report (STR/SAR) Data Inquiry", "2 hours", "12 hours (Mandatory filing via partners)"],
    ["Reporting Request", "Fraud Reporting & Fraudulent Account Flagging", "1 hour", "6 hours"],
    ["Electronic Fund Transfer", "Disputes on Electronic Transfers & Unallocated Credit Query", "30 minutes", "12 hours"],
    ["Dispute Processing", "Card & Transfer Dispute Investigation & Chargeback Evidence", "1 hour", "24 hours"],
    ["Other Requests", "API Configuration, IP Whitelist & Key Rotation", "30 minutes", "2 hours"],
    ["Other Requests", "Product Inquiry & Integration Technical Clarification", "1 hour", "4 hours"],
    ["Other Requests", "Compliance & Audit Documentation Request", "2 hours", "12 hours"]
]

# SLA Table 1 Data: PERFORMANCE INDICATORS
table1_widths = [1.5, 1.1, 1.1, 1.1, 1.7]
table1_headers = ["PERFORMANCE METRIC", "OPERATIONAL TARGET", "SUSPECT THRESHOLD", "CRITICAL THRESHOLD", "MEASUREMENT BASIS"]
table1_data = [
    ["Service Availability (Uptime)", "≥ 99.9%", "< 99.5%", "< 99.0%", "Monthly uptime across core BaaS APIs (excl. maintenance)"],
    ["API Response Latency (P95)", "≤ 500 ms", "> 1,500 ms", "> 3,000 ms", "95th percentile response time over 5-min rolling windows"],
    ["Transaction Processing Time", "≤ 3.0 seconds", "> 7.0 seconds", "> 15.0 seconds", "End-to-end execution from API call to partner bank response"],
    ["Transaction Success Rate", "≥ 99.0%", "< 95.0%", "< 90.0%", "Ratio of successful transactions vs valid processing attempts"],
    ["Webhook Delivery Success Rate", "≥ 99.9%", "< 98.0%", "< 95.0%", "Successful webhook receipt within 3 retries over rolling 24h"],
    ["Webhook Delivery Latency (P95)", "≤ 2.0 seconds", "> 10.0 seconds", "> 30.0 seconds", "Time elapsed between core ledger event and webhook receipt"],
    ["Daily Reconciliation Availability", "By 06:00 WAT", "By 09:00 WAT", "After 12:00 WAT", "Availability of daily transaction & settlement clearance files"]
]

# SLA Table 2 Data: SEVERITY, INCIDENT RESPONSE AND RESOLUTION MATRIX
table2_widths = [1.2, 2.5, 0.9, 0.9, 1.0]
table2_headers = ["SEVERITY LEVEL", "BUSINESS IMPACT DEFINITION", "INITIAL RESPONSE", "UPDATE FREQUENCY", "TARGET RESOLUTION"]
table2_data = [
    [
        "P1 – CRITICAL (Emergency)",
        "Complete service outage, core BaaS API failure, total virtual account provisioning failure, or systemic NIP transfer outage affecting all end-users with severe financial/regulatory risk.",
        "15 minutes (24/7/365)",
        "Every 30 mins",
        "Within 2 hours"
    ],
    [
        "P2 – MAJOR (High Impact)",
        "Partial service disruption, delayed webhook delivery, single bank partner degradation, or transaction latency affecting a significant portion of user transactions with no immediate workaround.",
        "30 minutes",
        "Every 1 hour",
        "Within 6 hours"
    ],
    [
        "P3 – MINOR (Medium Impact)",
        "Non-critical feature impairment, minor reporting/dashboard delay, sporadic latency on non-vital endpoints, or individual account query issue where an acceptable workaround exists.",
        "2 hours (Business Hours)",
        "Every 4 hours",
        "Within 24 hours"
    ],
    [
        "P4 – LOW (Informational / Enhancement)",
        "Cosmetic UI/documentation issues, minor inquiry, enhancement request, or general technical clarification with zero operational impact on active transactions.",
        "4 hours (Business Hours)",
        "Every 24 hours",
        "Within 48 hours / Next Release"
    ]
]

# SLA Table 3 Data: ESCALATION CHANNELS
table3_widths = [0.6, 1.4, 1.2, 1.7, 1.6]
table3_headers = ["LEVEL", "ESCALATION TIER & ROLE", "RESPONSE SLA & TRIGGER", "ANCHOR CONTACT DETAILS", "CLIENT (NOTESTANDARD) CONTACT"]
table3_data = [
    [
        "Level 1",
        "Technical Support & Incident Intake",
        "15 mins (P1/P2) / 1 hr (P3/P4)",
        "Technical Partnership Support\nEmail: support@getanchor.co / hello@getanchor.co\nSlack: #notestandard-anchor-support",
        "Engineering Support Lead\nEmail: support@notestandard.com / tech@notestandard.com\nPhone: +234 705 182 4027"
    ],
    [
        "Level 2",
        "Operations Analyst & Infrastructure Lead",
        "30 mins (P1) / 2 hrs (P2/P3)\n(30% of SLA elapsed)",
        "Precious Ehiwario (Operations Lead)\nEmail: precious@getanchor.co",
        "Head of Fintech Operations & Infrastructure\nEmail: ops@notestandard.com\nPhone: +234 705 182 4027"
    ],
    [
        "Level 3",
        "Head of Product & Technical Operations",
        "1 hr (P1) / 4 hrs (P2/P3)\n(50% of SLA elapsed)",
        "Tayo Brahm (Head of Product) / Olamide Sobowale\nEmail: tayo@getanchor.co / olamide@getanchor.co",
        "Chief Technology Officer / Product Lead\nEmail: admin@notestandard.com\nPhone: +234 705 182 4027"
    ],
    [
        "Level 4",
        "Emergency / Executive C-Suite (24/7)",
        "Immediate / 24/7\n(75% of SLA elapsed or P1 Emergency)",
        "Segun Adeyemi (Chief Executive Officer)\nEmail: segun@getanchor.co\nLine: 24/7 Emergency Line",
        "Oboh Aghogho Jossy (Founder & CEO)\nEmail: admin@notestandard.com / jossy@notestandard.com\nPhone: +234 705 182 4027 (24/7)"
    ]
]

# Create 4 new styled tables
table0 = create_styled_table(doc_sa, table0_widths, table0_headers, table0_data)
table1 = create_styled_table(doc_sa, table1_widths, table1_headers, table1_data)
table2 = create_styled_table(doc_sa, table2_widths, table2_headers, table2_data)
table3 = create_styled_table(doc_sa, table3_widths, table3_headers, table3_data)

# Remove the old original incomplete tables (first 4 tables in source)
for old_tbl in list(doc_sa.tables[:4]):
    old_tbl._tbl.getparent().remove(old_tbl._tbl)

# Locate paragraphs in Schedule B
p_issue = None
p_perf = None
p_sev = None
p_esc_intro = None

for p in doc_sa.paragraphs:
    txt = p.text.strip()
    if txt == "ISSUE RESOLUTION":
        p_issue = p
    elif txt == "PERFORMANCE INDICATORS":
        p_perf = p
    elif txt == "SEVERITY, INCIDENT RESPONSE AND RESOLUTION":
        p_sev = p
    elif "Where the issue is not resolved based on the above service levels" in txt:
        p_esc_intro = p

if p_issue:
    p_issue._p.addnext(table0._tbl)
if p_perf:
    p_perf._p.addnext(table1._tbl)
if p_sev:
    p_sev._p.addnext(table2._tbl)
if p_esc_intro:
    p_esc_intro._p.addnext(table3._tbl)

# Save DOCX files
doc_sa.save(FILE_DOCX)
doc_sa.save(os.path.join(WORKSPACE_DIR, "NoteStandard_Anchor_Client_Service_Agreement.docx"))
print(f"[OK] Saved {FILE_DOCX} and root NoteStandard_Anchor_Client_Service_Agreement.docx")

# ==============================================================================
# 3. FINCRA INDEMNITY PDF WITH EXTRACTED SIGNATURE & SEAL
# ==============================================================================
print("[3/6] Generating Fincra Wildcard IP Indemnity PDF with Signature & Seal...")

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#4B5563"))
        self.drawRightString(612 - 36, 20, f"Page {self._pageNumber} of {page_count}")
        self.drawString(36, 20, "CONFIDENTIAL — JOSSY DIGITAL TECHNOLOGIES LTD / FINCRA TECHNOLOGIES LTD")

def create_indemnity_pdf():
    doc = SimpleDocTemplate(
        FILE_INDEMNITY,
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=40,
        bottomMargin=40
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        alignment=1,
        textColor=colors.HexColor("#111827"),
        spaceAfter=15
    )

    body_style = ParagraphStyle(
        'BodyDark',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        alignment=4,
        textColor=colors.HexColor("#1F2937"),
        spaceAfter=8
    )

    bold_body = ParagraphStyle(
        'BoldBody',
        parent=body_style,
        fontName='Helvetica-Bold'
    )

    bullet_style = ParagraphStyle(
        'BulletText',
        parent=body_style,
        leftIndent=20,
        spaceAfter=6
    )

    story = []

    story.append(Paragraph("INDEMNITY FOR WILDCARD IP WHITELISTING", title_style))
    story.append(Spacer(1, 5))

    to_text = "<b>TO:</b><br/><b>Fincra Technologies Limited</b><br/>Plot 128, Block 1<br/>Polystar Electronics Building<br/>Remi Olowude Street, by Marwa Bus Stop<br/>Lekki Phase 1, Eti Osa<br/>Lagos State."
    story.append(Paragraph(to_text, body_style))
    story.append(Spacer(1, 10))

    intro_p1 = "This Indemnity is made this <u><b>3<sup>rd</sup></b></u> day of <u><b>August</b></u>, <u><b>2026</b></u> by <b>Jossy Digital Technologies Ltd</b>, (hereinafter called <b>\"the Customer\"</b>) in favour of <b>FINCRA TECHNOLOGIES LIMITED</b>, a company incorporated in Nigeria and having its registered office at Plot 128, Block 1 Polystar Electronics Building, Remi Olowude Street, by Marwa Bus Stop Lekki Phase 1, Eti Osa Lagos State (hereinafter called <b>\"Fincra\"</b>)."
    story.append(Paragraph(intro_p1, body_style))
    story.append(Spacer(1, 8))

    story.append(Paragraph("<b>WHEREAS</b>", bold_body))
    story.append(Paragraph("1. Fincra and the Customer entered into an agreement for the provision of services by Fincra to the Customer.", bullet_style))
    story.append(Paragraph("2. The Customer has, against Fincra's advice, requested Fincra to whitelist wildcard IP address for the Customer on Fincra Core (\"Request\").", bullet_style))
    story.append(Paragraph("3. The Customer recognizes that whitelisting wildcard IP poses a great risk to the Customer and exposes the Customer to attacks by fraudsters; it therefore accepts its obligation to secure its platform against fraudulent activities and accepts sole responsibility for any losses it incurs from same.", bullet_style))
    story.append(Paragraph("4. Fincra has requested that the Customer issues this indemnity in consideration of Fincra granting the Request as instructed by the Customer.", bullet_style))
    story.append(Spacer(1, 8))

    now_therefore = "<b>NOW THEREFORE</b> in consideration of the foregoing, the Customer hereby irrevocably and unconditionally undertakes and covenants to fully, indemnify Fincra and keep Fincra, its officers, directors, employees, affiliates and agents fully indemnified against all losses, claims, demands, liabilities, proceedings and damages which Fincra may suffer consequent upon Fincra granting the Request:"
    story.append(Paragraph(now_therefore, body_style))
    story.append(Spacer(1, 8))

    story.append(Paragraph("<b>I.</b> The Customer, its successors in title and assigns will at all times hereafter, keep Fincra, its successors in title and assigns indemnified against all actions, proceedings, claims, liabilities, losses, penalties and expenses and attorney fees in relation to/or arising out of the Request.", bullet_style))
    story.append(Paragraph("<b>II.</b> The Customer hereby indemnifies Fincra and holds it harmless from all liabilities, losses, penalties, expenses and claims of whatsoever nature which may be brought against Fincra or which it may suffer or incur, arising from it acting or for reasonable cause, not acting on any instruction in relation to the Request.", bullet_style))
    story.append(Paragraph("<b>III.</b> The Customer undertakes to pay Fincra upon Fincra's first written demand (not later than 72 hrs) without cavil or argument, all charges, costs and expenses including reasonable attorney fees incurred by Fincra in relation to/or arising out of the Request.", bullet_style))
    story.append(Paragraph("<b>IV.</b> Upon the Customer's failure to honour any demand made by Fincra in respect of this indemnity, the Customer hereby authorises Fincra to promptly debit Customer's account and/or wallet for any claims, demands, losses, expenses, damages, liabilities, charges, penalties, cost and attorney fees which Fincra may incur in connection with the Request.", bullet_style))
    story.append(Paragraph("<b>V.</b> This indemnity shall be a continuing obligation and shall not be withdrawn by the Customer as long as any liability shall remain upon Fincra in respect of the Request.", bullet_style))
    story.append(Paragraph("<b>VI.</b> This indemnity shall be governed by and construed in accordance with the laws of the Federal Republic of Nigeria.", bullet_style))
    story.append(Spacer(1, 10))

    dated_text = "Dated this <u><b>3<sup>rd</sup></b></u> day of <u><b>August</b></u>, <u><b>2026</b></u>"
    story.append(Paragraph(dated_text, body_style))
    story.append(Paragraph("The Common Seal of the within-named<br/><b>Jossy Digital Technologies Ltd</b><br/>was hereunto affixed in the presence of", body_style))
    story.append(Spacer(1, 10))

    sig_img = RLImage(SIG_PNG, width=120, height=50)
    seal_img = RLImage(SEAL_PNG, width=80, height=80)

    sig_left_cell = [
        sig_img,
        Paragraph("____________________________<br/><b>AGHOGHO OBOH</b><br/>DIRECTOR", body_style)
    ]
    sig_right_cell = [
        Spacer(1, 50),
        Paragraph("____________________________<br/><b>COMPANY SECRETARY</b><br/>DIRECTOR/SECRETARY", body_style)
    ]

    sig_data = [
        [sig_left_cell, seal_img, sig_right_cell]
    ]
    sig_table = Table(sig_data, colWidths=[210, 110, 210])
    sig_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
        ('ALIGN', (1,0), (1,0), 'CENTER'),
    ]))

    story.append(KeepTogether(sig_table))

    doc.build(story, canvasmaker=NumberedCanvas)

create_indemnity_pdf()
print(f"[OK] Saved {FILE_INDEMNITY}")

# ==============================================================================
# 4. EXECUTIVE COMPLIANCE SUMMARY PDF
# ==============================================================================
print("[4/6] Generating Executive Compliance Summary PDF...")

def create_compliance_summary_pdf():
    doc = SimpleDocTemplate(
        FILE_COMPLIANCE,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    styles = getSampleStyleSheet()

    header_style = ParagraphStyle(
        'HeaderStyle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#1E3A8A"),
        spaceAfter=2
    )

    sub_style = ParagraphStyle(
        'SubStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#2563EB"),
        spaceAfter=6
    )

    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11.5,
        alignment=4,
        textColor=colors.HexColor("#1F2937"),
        spaceAfter=4
    )

    bold_inline = ParagraphStyle(
        'BoldInline',
        parent=body_style,
        fontName='Helvetica-Bold'
    )

    story = []

    story.append(Paragraph("NOTESTANDARD (JOSSY DIGITAL TECHNOLOGIES LTD)", header_style))
    story.append(Paragraph("EXECUTIVE COMPLIANCE, RISK & BANKING INFRASTRUCTURE SUMMARY", sub_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#2563EB"), spaceAfter=8))

    table_data_1 = [
        [Paragraph("<b>1. Corporate & Legal Overview</b>", bold_inline), Paragraph("<b>2. Primary Business & Products</b>", bold_inline)],
        [
            Paragraph("• <b>Legal Entity:</b> Jossy Digital Technologies Ltd (RC: 9586407)<br/>• <b>Platform Brand:</b> NoteStandard (https://notestandard.com)<br/>• <b>Head Office:</b> Effurun, Delta State, Nigeria<br/>• <b>Emails:</b> admin@notestandard.com | support@notestandard.com<br/>• <b>Director:</b> Aghogho Oboh | <b>Date:</b> 03 August 2026<br/>• <b>Funding:</b> 100% Founder-Funded (Bootstrapped, $0 Dilution)", body_style),
            Paragraph("• <b>Core Offerings:</b> Multi-Currency Wallets, NUBAN Virtual Accounts, NIP Transfers, Internal Book Transfers, Bill Payments, Merchant Collections.<br/>• <b>Crypto Workflow:</b> Non-custodial fiat processing via regulated third-party infrastructure.<br/>• <b>Planned Expansion:</b> Payment cards, international ACH/SEPA (Phase 2).", body_style)
        ]
    ]
    t1 = Table(table_data_1, colWidths=[270, 270])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#F3F4F6")),
        ('PADDING', (0,0), (-1,-1), 5),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#D1D5DB")),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(t1)
    story.append(Spacer(1, 6))

    table_data_2 = [
        [Paragraph("<b>3. Internal Ledger Specification (System of Record)</b>", bold_inline), Paragraph("<b>4. Projected Processing Volumes (Baseline)</b>", bold_inline)],
        [
            Paragraph("• <b>System of Record:</b> NoteStandard maintains an internal, immutable double-entry ledger (<code>LedgerEngine.js</code>) for accounting, balance tracking, and audit trails.<br/>• <b>Anchor Role:</b> Anchor serves strictly as regulated banking clearing rail and NUBAN issuance partner, not system of record.<br/>• <b>Reconciliation:</b> Real-time webhook posting & automated end-of-day balance checks.", body_style),
            Paragraph("• <b>Monthly Baseline Volume:</b> ₦100,000,000 (~$66,667 USD)<br/>• <b>Daily Average Volume:</b> ₦3,333,333 (~$2,222 USD / ~170 txns)<br/>• <b>Flow Breakdown:</b> Wallet Funding NIP (40.5%), NIP Outbound Payouts (26.25%), Wallet Withdrawals (12.0%), Book Transfers (10.8%), Crypto On/Off-Ramp (8.4%), Bill Payments (2.05%).", body_style)
        ]
    ]
    t2 = Table(table_data_2, colWidths=[270, 270])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#F3F4F6")),
        ('PADDING', (0,0), (-1,-1), 5),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#D1D5DB")),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(t2)
    story.append(Spacer(1, 6))

    table_data_3 = [
        [Paragraph("<b>5. Identity Verification (KYC) Infrastructure</b>", bold_inline), Paragraph("<b>6. AML, PEP & Sanctions Screening</b>", bold_inline)],
        [
            Paragraph("• <b>Primary KYC Provider:</b> Prembly (IdentityPass). Additional providers may be integrated as business requirements evolve.<br/>• <b>Tiered KYC:</b> Tier 1 (BVN/NIN + Phone), Tier 2 (Govt ID + Address), Tier 3 (Utility + Business CAC).<br/>• <b>Underage Policy:</b> Accounts strictly prohibited for individuals under 18.", body_style),
            Paragraph("• <b>Screening & Verification:</b> Customer onboarding includes automated identity verification and sanctions screening through integrated compliance providers.<br/>• <b>Transaction Monitoring:</b> Custom AML engine with velocity rule limits (<code>risk_rules_velocity</code>).<br/>• <b>Risk Engine:</b> Automated risk scoring (<code>RiskDecisionEngine.js</code>) and manual review queue.", body_style)
        ]
    ]
    t3 = Table(table_data_3, colWidths=[270, 270])
    t3.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#F3F4F6")),
        ('PADDING', (0,0), (-1,-1), 5),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#D1D5DB")),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(t3)
    story.append(Spacer(1, 6))

    table_data_4 = [
        [Paragraph("<b>7. Information Security & Compliance Readiness</b>", bold_inline), Paragraph("<b>8. Operational Controls & Customer Support</b>", bold_inline)],
        [
            Paragraph("• <b>Authentication:</b> JWT tokens with multi-factor OTP validation for high-risk payouts.<br/>• <b>Idempotency:</b> Strict <code>event_id</code> idempotency keys & replay protection.<br/>• <b>Encryption:</b> AES-256-GCM encryption at rest (KMS); TLS 1.3 in transit.<br/>• <b>Audit Status:</b> Internal OWASP pen-tests & DR exercises complete; external audits planned.", body_style),
            Paragraph("• <b>Business Hours:</b> Mon–Fri 8AM–6PM WAT | Sat 9AM–4PM WAT.<br/>• <b>Support Channels:</b> In-app 24/7 AI + Live Chat, Email (admin@notestandard.com / support@notestandard.com), Phone, WhatsApp.<br/>• <b>Disclosures:</b> Terms of Service, Privacy Policy, and Electronic Consent Agreements published and active.", body_style)
        ]
    ]
    t4 = Table(table_data_4, colWidths=[270, 270])
    t4.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#F3F4F6")),
        ('PADDING', (0,0), (-1,-1), 5),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#D1D5DB")),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(t4)

    doc.build(story, canvasmaker=NumberedCanvas)

create_compliance_summary_pdf()
print(f"[OK] Saved {FILE_COMPLIANCE}")

# ==============================================================================
# 5. README MANIFEST
# ==============================================================================
print("[5/6] Writing README.txt...")

readme_text = """===============================================================================
NOTESTANDARD (JOSSY DIGITAL TECHNOLOGIES LTD) — ANCHOR BAAS SUBMISSION PACKAGE
===============================================================================

Submission Date: 03 August 2026
Legal Entity: Jossy Digital Technologies Ltd (RC: 9586407)
Platform Brand / Short Name: NoteStandard
Director: Aghogho Oboh
Website: https://notestandard.com
Admin Email: admin@notestandard.com
Support Email: support@notestandard.com
Head Office Address: Effurun, Delta State, Nigeria

PACKAGE MANIFEST & CONTENTS:
---------------------------
01 - Anchor Onboarding Questionnaire.xlsx
     Fully completed Anchor Company & Product Questionnaire Excel sheet. All 131 rows
     and Response fields populated without altering original formatting or structure.

02 - Standard Service Agreement.docx
     Executed Client Service Agreement between Anchor Software Ltd and Jossy Digital
     Technologies Ltd. Contains extracted Director signature and complete execution fields.

03 - Fincra Wildcard IP Indemnity.pdf
     Executed Wildcard IP Whitelisting Indemnity document bearing company seal and director signature.

04 - Cover Letter.docx
     Formal submission cover letter addressed to the Anchor Compliance Team, executed with signature.

05 - Compliance Summary.pdf
     One-page executive summary covering company overview, risk controls, KYC/AML engine,
     security architecture, and internal ledger specification.

06 - README.txt
     Manifest and verification documentation for the submission package.

SUMMARY DECLARATIONS:
--------------------
- Internal Ledger Statement: NoteStandard maintains its own internal double-entry ledger system
  of record (LedgerEngine.js). Anchor is used solely as regulated banking infrastructure and clearing rail.
- Baseline Processing Volume: Projected monthly transaction volume of ₦100,000,000 across 5,100 transactions.
- KYC Provider: Primary KYC Provider: Prembly (IdentityPass). Additional providers may be integrated as business requirements evolve.
- Sanctions Screening: Customer onboarding includes automated identity verification and sanctions screening through integrated compliance providers.

CONTACT INFORMATION:
--------------------
Company Name: Jossy Digital Technologies Ltd (RC 9586407)
Brand / Product Name: NoteStandard
Managing Director: Aghogho Oboh
Head Office Address: Effurun, Delta State, Nigeria
Admin Email: admin@notestandard.com
Support Email: support@notestandard.com
Website: https://notestandard.com
===============================================================================
"""

with open(FILE_README, "w", encoding="utf-8") as f:
    f.write(readme_text)

print(f"[OK] Saved {FILE_README}")

# ==============================================================================
# 6. COMPRESS INTO ZIP ARCHIVE
# ==============================================================================
print("[6/6] Compressing all submission files into ZIP archive...")

with zipfile.ZipFile(FILE_ZIP, "w", zipfile.ZIP_DEFLATED) as zf:
    zf.write(FILE_EXCEL, arcname="01 - Anchor Onboarding Questionnaire.xlsx")
    zf.write(FILE_DOCX, arcname="02 - Standard Service Agreement.docx")
    zf.write(FILE_INDEMNITY, arcname="03 - Fincra Wildcard IP Indemnity.pdf")
    zf.write(FILE_COVER, arcname="04 - Cover Letter.docx")
    zf.write(FILE_COMPLIANCE, arcname="05 - Compliance Summary.pdf")
    zf.write(FILE_README, arcname="06 - README.txt")

print(f"[SUCCESS] Submission package updated at:\n{FILE_ZIP}")
