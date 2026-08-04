import docx
import re

doc = docx.Document("temp_anchor_source/Updated Standard Service Agreement (2).docx")

print("--- AUDITING PARAGRAPHS FOR BLANK SPACES / DOTS / PLACEHOLDERS ---")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if any(k in text for k in ['…', '..', '____', 'COMPANY NAME', 'NAME:', 'DESIGNATION:', 'DATE:', 'THIS', 'made this', 'Client is a']):
        print(f"Para {i:3d}: \"{text}\"")

print("\n--- AUDITING TABLES FOR BLANK SPACES / DOTS / PLACEHOLDERS ---")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if any(k in text for k in ['…', '..', '____', 'COMPANY NAME', 'NAME:', 'DESIGNATION:', 'DATE:']):
                print(f"Table {t_idx} R{r_idx} C{c_idx}: \"{text}\"")
