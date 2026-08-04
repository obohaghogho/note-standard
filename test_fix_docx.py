import docx
from docx.shared import Inches

src_docx = "temp_anchor_source/Updated Standard Service Agreement (2).docx"
dst_docx = "test_fixed_agreement.docx"
sig_png = "temp_anchor_source/signature_transparent.png"

doc = docx.Document(src_docx)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    # Para 12: Client Name on title page
    if i == 12 or "………………………………" in text or (len(text) > 5 and all(c in "…." for c in text)):
        p.text = "JOSSY DIGITAL TECHNOLOGIES LTD"
    
    # Para 16: Date on title page
    elif i == 16 or "THIS" in text and "2026" in text and ("…" in text or ".." in text):
        p.text = "THIS 03 DAY OF AUGUST 2026"
    
    # Para 21: Date in opening line
    elif i == 21 or ("made this" in text and "2026" in text):
        p.text = "This Client Service Agreement (“Agreement”) is made this 03 day of August 2026."
    
    # Para 25: Client legal details
    elif i == 25 or ("assigns and successors" in text and "Client" in text and ("…" in text or ".." in text)):
        p.text = "JOSSY DIGITAL TECHNOLOGIES LTD, a limited liability company duly incorporated under the laws of the Federal Republic of Nigeria, with RC Number: 9586407 and having its registered address at Effurun, Delta State, Nigeria (where the context so admits includes its assigns and successors) (“the Client”)."
    
    # Para 28: Client business description
    elif i == 28 or "The Client is a…" in text or "The Client is a." in text:
        p.text = "The Client is a fintech and digital workspace platform providing multi-currency wallets, virtual accounts, NIP bank transfers, internal ledger transfers, merchant collections, bill payments, and treasury technology services."

    # Anchor Date
    elif "DATE:" in text and i < 377 and "OLUWASEGUN" in doc.paragraphs[i-2].text:
        p.text = "DATE: 03 August 2026"

    # Client Execution block
    elif "FOR: COMPANY NAME [CLIENT]" in text or "FOR: COMPANY NAME" in text:
        p.text = "FOR: JOSSY DIGITAL TECHNOLOGIES LTD [CLIENT]"
        # Update subsequent paragraphs in execution block
        if i + 2 < len(doc.paragraphs) and "SIGNATURE:" in doc.paragraphs[i+2].text:
            sig_p = doc.paragraphs[i+2]
            sig_p.text = "SIGNATURE: "
            run = sig_p.add_run()
            run.add_picture(sig_png, width=Inches(1.8))
        if i + 3 < len(doc.paragraphs):
            doc.paragraphs[i+3].text = "NAME: AGHOGHO OBOH"
        if i + 4 < len(doc.paragraphs):
            doc.paragraphs[i+4].text = "DESIGNATION: DIRECTOR"
        if i + 5 < len(doc.paragraphs):
            doc.paragraphs[i+5].text = "DATE: 03 August 2026"

    # Schedules
    elif "Please find attached a list of our high risk and prohibited Customers" in text:
        p.text = "NoteStandard prohibits onboarding high-risk/restricted entities including sanctions-listed persons, un-hosted gambling operators, adult content services, and unregulated darknet vendors."
    elif "Please find attached a list of our high risk and prohibited Jurisdictions" in text:
        p.text = "NoteStandard enforces FATF and OFAC high-risk and non-cooperative jurisdiction blacklists including North Korea (DPRK), Iran, Myanmar, Syria, and Cuba."

doc.save(dst_docx)
print("Saved test_fixed_agreement.docx successfully!")
