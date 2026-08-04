import docx

doc = docx.Document("temp_anchor_source/Updated Standard Service Agreement (2).docx")
for i in range(365, len(doc.paragraphs)):
    print(f"Para {i:3d}: '{doc.paragraphs[i].text}'")
